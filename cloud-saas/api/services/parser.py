"""
AeroSync Cloud - 文档解析服务
支持：Excel (.xlsx/.xls)、PDF (.pdf)、Word (.doc/.docx)
从 OSS 下载到临时文件后解析
"""
import os
import re
import json
import tempfile
from typing import Dict, List, Any

import oss2
from openpyxl import load_workbook
import pdfplumber
from docx import Document

from api.core.config import settings
from api.core.logging_config import get_logger

logger = get_logger("parser")


def get_oss_bucket():
    """获取阿里云 OSS Bucket 实例"""
    auth = oss2.Auth(settings.OSS_ACCESS_KEY_ID, settings.OSS_ACCESS_KEY_SECRET)
    bucket = oss2.Bucket(auth, settings.OSS_ENDPOINT, settings.OSS_BUCKET)
    return bucket


def download_from_oss(object_key: str, local_path: str) -> bool:
    """从 OSS 下载文件到本地临时路径"""
    # 检查 OSS 配置是否完整
    if not all([
        settings.OSS_ACCESS_KEY_ID,
        settings.OSS_ACCESS_KEY_SECRET,
        settings.OSS_ENDPOINT,
        settings.OSS_BUCKET
    ]):
        logger.warning(f"OSS 凭证未配置，跳过下载: {object_key}")
        return False

    try:
        bucket = get_oss_bucket()
        bucket.get_object_to_file(object_key, local_path)
        logger.info(f"OSS 下载成功: {object_key}")
        return True
    except Exception as e:
        logger.error(f"从OSS下载失败: {object_key}, 错误: {e}")
        return False


def parse_document(object_key: str, filename: str) -> Dict[str, Any]:
    """
    主解析入口：根据文件类型分发到对应解析器
    若启用混合解析模式 (ENABLE_HYBRID_PARSER=1)，则使用 Docling + 基础解析器混合方案
    Returns: {"type": "excel|pdf|word", "text": "全文文本", "structured": {...}, "tables": [...]}
    """
    suffix = os.path.splitext(filename)[1].lower()

    # 创建临时文件
    fd, temp_path = tempfile.mkstemp(suffix=suffix)
    os.close(fd)

    try:
        # 1. 从 OSS 下载
        downloaded = download_from_oss(object_key, temp_path)

        if not downloaded:
            # OSS 未配置或下载失败，返回模拟/空数据
            logger.warning(f"OSS 下载失败或未配置，返回空解析数据: {filename}")
            return _empty_parse_result(filename, suffix)

        # 2. 启用混合解析模式（Phase 2）
        if settings.ENABLE_HYBRID_PARSER:
            try:
                from api.services.hybrid_parser import parse_with_hybrid
                logger.info(f"使用混合解析模式: {filename}")
                return parse_with_hybrid(temp_path, filename)
            except Exception as e:
                logger.warning(f"混合解析失败，回退到基础解析: {e}")

        # 3. 按类型解析（基础模式）
        if suffix in ['.xlsx', '.xls']:
            return parse_excel(temp_path, filename)
        elif suffix == '.pdf':
            return parse_pdf(temp_path, filename)
        elif suffix in ['.doc', '.docx']:
            return parse_word(temp_path, filename)
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")

    finally:
        # 清理临时文件
        if os.path.exists(temp_path):
            os.unlink(temp_path)


def _empty_parse_result(filename: str, suffix: str) -> Dict[str, Any]:
    """
    当 OSS 未配置或下载失败时，返回一个空但结构完整的解析结果
    """
    if suffix in ['.xlsx', '.xls']:
        doc_type = "excel"
    elif suffix == '.pdf':
        doc_type = "pdf"
    elif suffix in ['.doc', '.docx']:
        doc_type = "word"
    else:
        doc_type = "unknown"

    return {
        "type": doc_type,
        "filename": filename,
        "text": "",
        "structured": {
            "note": "OSS 未配置或下载失败，无法获取文件内容",
            "sheet_count": 0,
            "sheet_names": [],
            "headers": [],
            "total_rows": 0,
            "page_count": 0,
            "extracted_tables": 0,
            "paragraph_count": 0,
            "table_count": 0,
        },
        "tables": []
    }


def parse_excel(file_path: str, filename: str) -> Dict[str, Any]:
    """解析 Excel 文件：提取所有 sheet 的文本和表格数据"""
    wb = load_workbook(file_path, data_only=True)
    all_text_lines = []
    all_tables = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_text_lines.append(f"=== Sheet: {sheet_name} ===")

        sheet_data = []
        for row in ws.iter_rows(values_only=True):
            # 过滤全空行
            if any(cell is not None for cell in row):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                sheet_data.append(row_data)
                all_text_lines.append("\t".join(row_data))

        all_tables.append({
            "sheet": sheet_name,
            "rows": sheet_data
        })

    full_text = "\n".join(all_text_lines)

    # 智能识别表头（取第一个 sheet 的第一行）
    headers = all_tables[0]["rows"][0] if all_tables and all_tables[0]["rows"] else []

    return {
        "type": "excel",
        "filename": filename,
        "text": full_text,
        "structured": {
            "sheet_count": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
            "headers": headers,
            "total_rows": sum(len(t["rows"]) for t in all_tables),
        },
        "tables": all_tables
    }


def parse_pdf(file_path: str, filename: str) -> Dict[str, Any]:
    """解析 PDF 文件：提取文本和表格"""
    all_text_lines = []
    all_tables = []
    page_count = 0

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            # 提取文本
            text = page.extract_text()
            if text:
                all_text_lines.append(f"=== Page {i + 1} ===")
                all_text_lines.append(text)

            # 提取表格
            tables = page.extract_tables()
            for j, table in enumerate(tables):
                if table:
                    rows = [[str(cell) if cell else "" for cell in row] for row in table]
                    all_tables.append({
                        "page": i + 1,
                        "index": j,
                        "rows": rows
                    })

    full_text = "\n".join(all_text_lines)

    return {
        "type": "pdf",
        "filename": filename,
        "text": full_text,
        "structured": {
            "page_count": page_count,
            "extracted_tables": len(all_tables),
        },
        "tables": all_tables
    }


def parse_word(file_path: str, filename: str) -> Dict[str, Any]:
    """解析 Word 文件：提取段落文本和表格"""
    doc = Document(file_path)

    # 提取段落
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # 提取表格
    all_tables = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            rows.append(row_data)
        all_tables.append({
            "index": i,
            "rows": rows
        })

    full_text = "\n".join(paragraphs)

    return {
        "type": "word",
        "filename": filename,
        "text": full_text,
        "structured": {
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
        },
        "tables": all_tables
    }


# ==================== 航空领域专用提取器 ====================

# 件号 (Part Number) 正则模式
PART_NUMBER_PATTERNS = [
    r'(?:件号|件號|P/N|Part\s*Number|PART\s*NO)[：:\s]*([A-Z0-9\-]{5,20})',
    r'\b([A-Z]{2}\d{3,}[A-Z0-9\-]*)\b',  # 如: BACB30FM8A4
    r'\b(\d{2}[A-Z]\d{4,}[A-Z0-9\-]*)\b',  # 如: 10-60516-14
]

# 机型正则模式
AIRCRAFT_MODEL_PATTERNS = [
    r'(B\d{3}-[A-Z0-9]+)',  # 波音: B737-800
    r'(A\d{3}[A-Z]?)',       # 空客: A320, A350
    r'(C\d{3})',             # C919
    r'(E\d{3})',             # 巴航: E190
]

# ATA 章节模式
ATA_PATTERN = r'ATA\s*(\d{2,3})'

# 数量模式
QUANTITY_PATTERNS = [
    r'(?:数量|數量|QTY|Quantity)[：:\s]*(\d+)',
    r'\b(\d+)\s*(?:EA|ea|件|个|PC|pc)\b',
]


def extract_aviation_entities(text: str) -> Dict[str, Any]:
    """
    从文本中提取航空领域关键实体
    返回: {"part_numbers": [...], "aircraft_models": [...], "ata_chapters": [...], "quantities": [...]}
    """
    entities = {
        "part_numbers": [],
        "aircraft_models": [],
        "ata_chapters": [],
        "quantities": [],
        "pma_marked": "PMA" in text.upper(),
        "landing_gear_marked": bool(re.search(r'起落架|landing\s*gear', text, re.IGNORECASE)),
    }

    # 提取件号
    for pattern in PART_NUMBER_PATTERNS:
        matches = re.findall(pattern, text)
        entities["part_numbers"].extend(matches)
    entities["part_numbers"] = list(set(entities["part_numbers"]))[:20]  # 去重限流

    # 提取机型
    for pattern in AIRCRAFT_MODEL_PATTERNS:
        matches = re.findall(pattern, text)
        entities["aircraft_models"].extend(matches)
    entities["aircraft_models"] = list(set(entities["aircraft_models"]))

    # 提取 ATA
    ata_matches = re.findall(ATA_PATTERN, text)
    entities["ata_chapters"] = list(set(ata_matches))

    # 提取数量
    for pattern in QUANTITY_PATTERNS:
        matches = re.findall(pattern, text)
        entities["quantities"].extend(matches)
    entities["quantities"] = list(set(entities["quantities"]))

    return entities
